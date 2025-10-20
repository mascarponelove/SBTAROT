import random
from pathlib import Path
from docx import Document

class TarotDeck:
    """
    Enhanced Tarot Deck with file system integration.
    Handles images and contextual meanings from Word documents.
    Deck is randomized on creation for robust shuffling.
    """
    
    def __init__(self, assets_path="assets"):
        self.assets_path = Path(assets_path)
        
        # Major Arcana card names (standardized)
        self.major_names = [
            "FOOL", "MAGICIAN", "HIGH_PRIESTESS", "EMPRESS", 
            "EMPEROR", "HIEROPHANT", "LOVERS", "CHARIOT", 
            "STRENGTH", "HERMIT", "WHEEL_OF_FORTUNE", "JUSTICE", 
            "HANGED_MAN", "DEATH", "TEMPERANCE", "DEVIL", 
            "TOWER", "STAR", "MOON", "SUN", "JUDGEMENT", "WORLD"
        ]
        
        # Minor Arcana parameters
        self.suits = ["cups", "pentacles", "swords", "wands"]
        self.ranks = ["ACE", "TWO", "THREE", "FOUR", "FIVE", "SIX", 
                      "SEVEN", "EIGHT", "NINE", "TEN", "PAGE", 
                      "KNIGHT", "QUEEN", "KING"]
        
        # Build full deck structure
        self.deck = self._build_deck()
        
        # CRITICAL: Randomize the base deck order immediately
        random.shuffle(self.deck)
        
        # Current deck starts as a copy of the randomized deck
        self.current_deck = self.deck.copy()
    
    def _build_deck(self):
        """Construct deck with file path references."""
        deck = []
        
        # Add Major Arcana
        for card_name in self.major_names:
            deck.append({
                "type": "major",
                "name": card_name,
                "display_name": card_name.replace("_", " ").title(),
                "image_path": f"assets/images/major/{card_name}.png",
                "meaning_path": f"assets/meanings/major/{card_name}.docx"
            })
        
        # Add Minor Arcana
        for suit in self.suits:
            for rank in self.ranks:
                deck.append({
                    "type": "minor",
                    "suit": suit,
                    "rank": rank,
                    "name": f"{rank}_OF_{suit.upper()}",
                    "display_name": f"{rank.title()} of {suit.title()}",
                    "image_path": f"assets/images/minor/{suit}/{rank}.png",
                    "meaning_path": f"assets/meanings/minor/{suit}/{rank}.docx"
                })
        
        return deck
    
    def shuffle(self):
        """
        Fisher-Yates shuffle implementation.
        Resets to full deck and randomizes order.
        """
        self.current_deck = self.deck.copy()
        random.shuffle(self.current_deck)
        return {"status": "shuffled", "cards_remaining": len(self.current_deck)}
    
    def draw_card(self):
        """
        Draw one card from deck.
        Card is removed and cannot be drawn again until reset/shuffle.
        """
        if not self.current_deck:
            return None
        return self.current_deck.pop()
    
    def extract_meaning(self, card, context="Soul"):
        """
        Extract meaning from Word document based on context.
        Reads 2-column table format.
        
        Args:
            card (dict): Card dictionary from deck
            context (str): Context type (Soul, Essence/Prediction, Past, Present, 
                          Future, Health, Profession, Relationship, Guidance)
        
        Returns:
            str: Contextual meaning
        """
        meaning_file = self.assets_path / card["meaning_path"].replace("assets/", "")
        
        if not meaning_file.exists():
            return f"⚠️ Meaning file not found: {card['display_name']}. Please add the Word document at: {card['meaning_path']}"
        
        try:
            doc = Document(meaning_file)
            
            # Parse 2-column table
            if doc.tables:
                table = doc.tables[0]
                
                # Store all card data
                card_data = {}
                
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        category = cells[0].text.strip()
                        content = cells[1].text.strip()
                        
                        if category and content:
                            card_data[category] = content
                
                # Return requested context (exact match first)
                if context in card_data:
                    return card_data[context]
                
                # Try case-insensitive match
                for key in card_data.keys():
                    if key.lower() == context.lower():
                        return card_data[key]
                
                # Try partial match
                for key in card_data.keys():
                    if context.lower() in key.lower() or key.lower() in context.lower():
                        return card_data[key]
                
                # If still not found, list available categories
                available = ", ".join(card_data.keys())
                return f"Context '{context}' not found. Available categories: {available}"
            
            return "No table found in Word document. Please ensure your .docx file contains a table."
        
        except Exception as e:
            return f"Error reading meaning file: {str(e)}"
    
    def get_all_card_data(self, card):
        """
        Get all data from card's Word document (for displaying Yes/No, +/-, etc.)
        
        Args:
            card (dict): Card dictionary from deck
        
        Returns:
            dict: All card data from Word document
        """
        meaning_file = self.assets_path / card["meaning_path"].replace("assets/", "")
        
        if not meaning_file.exists():
            return {}
        
        try:
            doc = Document(meaning_file)
            
            if doc.tables:
                table = doc.tables[0]
                card_data = {}
                
                for row in table.rows:
                    cells = row.cells
                    if len(cells) >= 2:
                        category = cells[0].text.strip()
                        content = cells[1].text.strip()
                        
                        if category and content:
                            card_data[category] = content
                
                return card_data
            
            return {}
        
        except Exception as e:
            return {"error": str(e)}
    
    def reset(self):
        """
        Reset deck to full 78 cards with new random order.
        This ensures every reset gives a fresh random arrangement.
        """
        self.current_deck = self.deck.copy()
        random.shuffle(self.current_deck)
        return {"status": "reset", "cards_remaining": len(self.current_deck)}
